import os, io, json
from datetime import datetime, date
from flask import (
    Flask, request, jsonify, redirect, url_for, flash, render_template,
    send_file
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadSignature
from openai import OpenAI
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pt_templates import PT_TEMPLATES, OT_TEMPLATES, pt_parse_template, ot_parse_template
from models import db, Therapist, Patient, PTNote, get_pt_note_for_patient


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev_secret_key_change_me")
db_path = os.path.join(os.getcwd(), 'db.sqlite3')
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{db_path}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False


db.init_app(app)
with app.app_context():
    try:
        db.engine.execute('ALTER TABLE ptnote ADD COLUMN fields_json TEXT')
        print("Migration succeeded!")
    except Exception as e:
        print("Migration failed or column already exists:", e)
        
# ==================================== OPENAI SETUP ====================================
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
client = OpenAI(api_key=OPENAI_API_KEY)
MODEL = "gpt-4o-mini"

# ==================================== LOGIN MANAGER ====================================
login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return Therapist.query.get(int(user_id))

# =========================== CREATE TABLES & ADMIN IF NONE ============================
with app.app_context():
    db.create_all()
    if not Therapist.query.filter_by(username="admin").first():
        hashed_pw = generate_password_hash("admin123")
        admin = Therapist(
            username="admin",
            password=hashed_pw,
            first_name="Admin",
            last_name="User",
            credentials="PT",
            email="admin@example.com",
            phone="555-555-5555",
        )
        db.session.add(admin)
        db.session.commit()
        print("Default admin user created: username=admin, password=admin123")

# ==================================== INDEX REDIRECT ====================================
@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

# ==================================== LOGIN ====================================
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Therapist.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash("Invalid username or password", "danger")
    return render_template('login.html')

# ==================================== LOGOUT ====================================
@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

# ==================================== DASHBOARD ====================================
@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')

# ==================================== REGISTER ====================================
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        email = request.form['email']

        if Therapist.query.filter_by(username=username).first():
            flash("Username already exists", "danger")
            return render_template('register.html')
        if Therapist.query.filter_by(email=email).first():
            flash("Email already registered", "danger")
            return render_template('register.html')

        hashed_pw = generate_password_hash(password)
        therapist = Therapist(
            username=username,
            password=hashed_pw,
            email=email,
        )
        db.session.add(therapist)
        db.session.commit()
        flash("Registration successful! Please log in.", "success")
        return redirect(url_for('login'))
    return render_template('register.html')

# ==================================== FORGOT PASSWORD ====================================

s = URLSafeTimedSerializer(app.secret_key)

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email'].strip().lower()
        user = Therapist.query.filter_by(email=email).first()
        if user:
            token = s.dumps(user.id, salt='password-reset')
            reset_url = url_for('reset_password', token=token, _external=True)
            print(f"Reset link for {email}: {reset_url}")  # Send by email in production
        flash("If this email exists, a reset link has been sent.", "success")
        return redirect(url_for('forgot_password'))
    return render_template('forgot_password.html')
    
# ==================================== RESET PASSWORD ====================================

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        user_id = s.loads(token, salt='password-reset', max_age=3600)
    except SignatureExpired:
        flash("Reset link expired. Please try again.", "danger")
        return redirect(url_for('forgot_password'))
    except BadSignature:
        flash("Invalid or tampered link.", "danger")
        return redirect(url_for('forgot_password'))

    user = Therapist.query.get(user_id)
    if not user:
        flash("User not found.", "danger")
        return redirect(url_for('forgot_password'))

    if request.method == 'POST':
        password = request.form['password']
        confirm = request.form['confirm']
        if not password or password != confirm:
            flash("Passwords do not match.", "danger")
            return render_template('reset_password.html')
        user.password = generate_password_hash(password)
        db.session.commit()
        flash("Password reset successful. You can log in now.", "success")
        return redirect(url_for('login'))
    return render_template('reset_password.html')
    
# ==================================== ADD PATIENTS ====================================
@app.route('/patients', methods=['GET'])
@login_required
def view_patients():
    query = request.args.get('q', '').strip().lower()
    if query:
        patients = Patient.query.filter(Patient.name.ilike(f"%{query}%")).order_by(Patient.name.asc()).all()
    else:
        patients = Patient.query.order_by(Patient.name.asc()).all()
    return render_template('patients.html', patients=patients, query=query)
    
@app.route('/add_patient', methods=['GET', 'POST'])
def add_patient():
    if request.method == 'POST':
        name = request.form['name']
        dob_str = request.form['dob']
        gender = request.form['gender']

        try:
            dob = datetime.strptime(dob_str, '%m/%d/%Y').date()
        except ValueError:
            error_message = "Invalid date format for DOB. Use MM/DD/YYYY."
            return render_template('add_patient.html', error_message=error_message)

        new_patient = Patient(name=name, dob=dob, gender=gender)
        db.session.add(new_patient)
        db.session.commit()

        flash('Patient added successfully!', 'success')
        return redirect(url_for('view_patients'))

    # For GET request, render the form
    return render_template('add_patient.html')




#==================================== EDIT PATIENTS ====================================

@app.route('/edit_patient/<int:patient_id>', methods=['GET', 'POST'])
@login_required
def edit_patient(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    if request.method == 'POST':
        # Update fields from form data
        patient.name = request.form['name']
        patient.dob = request.form['dob']
        patient.gender = request.form['gender']
        # ... other fields
        db.session.commit()
        flash("Patient info updated.", "success")
        return redirect(url_for('view_patients'))
    return render_template('edit_patient.html', patient=patient)
    
#==================================== DELETE PATIENTS  ====================================

@app.route('/delete_patient/<int:patient_id>', methods=['POST'])
@login_required
def delete_patient(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    # Also delete notes if you want to cascade delete
    db.session.delete(patient)
    db.session.commit()
    flash("Patient deleted.", "success")
    return redirect(url_for('view_patients'))

#==================================== VIEW NOTES  ====================================

@app.route('/view_pt_notes/<int:patient_id>')
@login_required
def view_pt_notes(patient_id):
    print("Fetching notes for patient_id:", patient_id, type(patient_id))
    notes = PTNote.query.filter_by(patient_id=patient_id).order_by(PTNote.created_at.desc()).all()
    print("Fetched notes:", [(n.id, n.patient_id) for n in notes])
    return render_template('pt_notes.html', notes=notes)


@app.route('/view_note/<int:note_id>')
@login_required
def view_note(note_id):
    note = PTNote.query.get_or_404(note_id)
    fields = json.loads(note.fields_json) if note.fields_json else {}
    return render_template('view_note.html', fields=fields, note=note)


@app.route('/patients/<int:patient_id>/ot-notes')
@login_required
def view_ot_notes(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    return render_template('ot_notes.html', patient=patient)

#==================================== EDIT PT NOTES ====================================
@app.route('/pt_notes/<int:note_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_pt_note(note_id):
    note = PTNote.query.get_or_404(note_id)
    if request.method == 'POST':
        note.content = request.form['content']
        note.doc_type = request.form.get('doc_type', note.doc_type)
        db.session.commit()
        flash('Note updated!')
        return redirect(url_for('pt_notes'))
    return render_template('edit_pt_note.html', note=note)

#==================================== DELETE PT NOTES ====================================

@app.route('/pt_notes/<int:note_id>/delete', methods=['POST'])
@login_required
def delete_pt_note(note_id):
    note = PTNote.query.get_or_404(note_id)
    db.session.delete(note)
    db.session.commit()
    flash('Note deleted!')
    return redirect(url_for('pt_notes'))

#==================================== LOAD PT NOTE ====================================

@app.route('/pt_notes/<int:note_id>/load')
@login_required
def load_pt_note(note_id):
    note = PTNote.query.get_or_404(note_id)
    try:
        loaded_fields = json.loads(note.fields_json or '{}')
    except Exception:
        loaded_fields = {}

    loaded_doc_type = note.doc_type if note.doc_type is not None else None

    return render_template('pt_eval.html', loaded_note=loaded_fields, loaded_doc_type=loaded_doc_type)



#==================================== PATIENT SEARCH ====================================

@app.route('/patients', methods=['GET'])
@login_required
def view_single_patient():
    query = request.args.get('q', '').strip().lower()

    if query:
        patients = Patient.query.filter(Patient.name.ilike(f"%{query}%")).order_by(
            db.func.split_part(Patient.name, ' ', -1), Patient.name
        ).all()
    else:
        patients = Patient.query.order_by(
            db.func.split_part(Patient.name, ' ', -1), Patient.name
        ).all()

    return render_template('patients.html', patients=patients, query=query)

@app.route('/api/patients')
@login_required
def api_patients():
    query = request.args.get('q', '').strip()
    results = Patient.query
    if query:
        results = results.filter(Patient.name.ilike(f"%{query}%"))
    results = results.order_by(Patient.name.asc()).all()
    return jsonify([{"id": p.id, "name": p.name} for p in results])
    
# #==================================== LOAD PT TEMPLATES ====================================
@app.route("/pt_load_template", methods=["POST"])
@login_required
def pt_load_template():
    data = request.get_json()
    template_name = data.get("template", "")
    if not template_name:
        return jsonify(list(PT_TEMPLATES.keys()))
    else:
        return jsonify(PT_TEMPLATES.get(template_name, {}))

@app.route("/ot_load_template", methods=["POST"])
@login_required
def ot_load_template():
    req = request.get_json() or {}
    template = req.get("template", "")
    if template:
        return jsonify(OT_TEMPLATES.get(template, {}))
    else:
        return jsonify(list(OT_TEMPLATES.keys()))
        


#==================================== PT RENDER EVAL ====================================

@app.route('/pt-eval')
@login_required
def pt_eval():
    # your PT Eval page rendering logic here
    loaded_note = None
    # fetch loaded_note if needed
    return render_template('pt_eval.html', loaded_note=loaded_note)
        
    
#========================== AI GENERATE PT DIFFERENTIAL DX ==============================

@app.route("/pt_generate_diffdx", methods=["POST"])
@login_required
def pt_generate_diffdx():
    f = request.json.get("fields", {})
    pain = "; ".join(f"{lbl}: {f.get(key,'')}"
                      for lbl, key in [
                          ("Area/Location", "pain_location"),
                          ("Onset", "pain_onset"),
                          ("Condition", "pain_condition"),
                          ("Mechanism", "pain_mechanism"),
                          ("Rating", "pain_rating"),
                          ("Frequency", "pain_frequency"),
                          ("Description", "pain_description"),
                          ("Aggravating", "pain_aggravating"),
                          ("Relieved", "pain_relieved"),
                          ("Interferes", "pain_interferes"),
                      ])
    prompt = (
        "You are a PT clinical assistant. Based on the following evaluation details, "
        "provide a concise statement of the most clinically-associated PT differential diagnosis. "
        "Do NOT state as fact or as a medical diagnosis—use only language such as 'symptoms and clinical findings are associated with or consistent with' the diagnosis. "
        "Keep the statement clean and PT-relevant:\n\n"
        f"Subjective:\n{f.get('subjective','')}\n\n"
        f"Pain:\n{pain}\n\n"
        f"Objective:\nPosture: {f.get('posture','')}\n"
        f"ROM: {f.get('rom','')}\n"
        f"Strength: {f.get('strength','')}\n"
    )
    result = gpt_call(prompt, max_tokens=250)
    return jsonify({"result": result})
    
#==================================== AI GENERATE PT SUMMARY ====================================

@app.route("/pt_generate_summary", methods=["POST"])
@login_required
def pt_generate_summary():
    f = request.json.get("fields", {})
    name = (
        f.get("name")
        or f.get("pt_patient_name")
        or f.get("patient_name")
        or f.get("full_name")
        or "Pt"
    )
    dob = f.get("dob")
    age = "X"
    if dob:
        for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m-%d-%Y"):
            try:
                dob_dt = datetime.strptime(dob, fmt)
                today_dt = date.today()
                age = today_dt.year - dob_dt.year - ((today_dt.month, today_dt.day) < (dob_dt.month, dob_dt.day))
                break
            except Exception:
                continue
    else:
        age = f.get("age", "X")

    gender = f.get("gender", "patient").lower()
    pmh = f.get("history", "no significant history")
    today = f.get("currentdate", date.today().strftime("%m/%d/%Y"))
    subj = f.get("subjective", "")
    moi = f.get("pain_mechanism", "")
    meddiag = f.get("meddiag", "") or f.get("medical_diagnosis", "")
    dx = f.get("diffdx", "")
    strg = f.get("strength", "")
    rom = f.get("rom", "")
    impair = f.get("impairments", "")
    func = f.get("functional", "")

    prompt = (
        "Generate a concise, 7-8 sentence Physical Therapy assessment summary that is Medicare compliant for PT documentation. "
        "Use only abbreviations (e.g., HEP, ADLs, LBP, STM, TherEx) and NEVER spell out abbreviations. "
        "Never use 'the patient'; use 'Pt' as the subject. "
        "Do NOT use parentheses, asterisks, or markdown formatting in your response. "
        "Do NOT use 'Diagnosis:' as a label—refer directly to the diagnosis in clinical sentences. "
        "Do NOT state or conclude a medical diagnosis—use clinical phrasing such as 'symptoms and clinical findings are associated with' the medical diagnosis and PT clinical impression. "
        f"Start with: \"{name}, a {age} y/o {gender} with relevant history of {pmh}.\" "
        f"Include: PT initial eval on {today} for {subj}. "
        f"If available, mention the mechanism of injury: {moi}. "
        f"State: Pt has symptoms and clinical findings associated with the referring medical diagnosis of {meddiag}. Clinical findings are consistent with PT differential diagnosis of {dx} based on assessment. "
        f"Summarize current impairments (strength: {strg}; ROM: {rom}; balance/mobility: {impair}). "
        f"Summarize functional/activity limitations: {func}. "
        "End with a professional prognosis stating that skilled PT is medically necessary to address impairments and support return to PLOF. "
        "Do NOT use bulleted or numbered lists—compose a single, well-written summary paragraph."
    )
    result = gpt_call(prompt, max_tokens=500)
    return jsonify({"result": result})
    
#==================================== AI GENERATE PT GOALS ====================================
@app.route('/pt_generate_goals', methods=['POST'])
@login_required
def pt_generate_goals():
    fields = request.json.get("fields", {})

    summary = fields.get("summary", "")
    strength = fields.get("strength", "")
    rom = fields.get("rom", "")
    impairments = fields.get("impairments", "")
    functional = fields.get("functional", "")
    meddiag = fields.get("meddiag", "")
    pain_location = fields.get("pain_location", "")

    prompt = f"""
You are a clinical assistant helping a PT write documentation.
Below is a summary of the patient's evaluation and findings:
Diagnosis/Region: {meddiag or pain_location}
Summary: {summary}
Strength: {strength}
ROM: {rom}
Impairments: {impairments}
Functional Limitations: {functional}

Using ONLY the above provided eval info, generate clinically-appropriate, Medicare-compliant short-term and long-term PT goals for the region/problem described.
Each goal must be functionally focused and follow this EXACT format, time frame, and language example (do NOT copy example content, use it as a style guide):

Short-Term Goals (1–12 visits):
1. Pt will report [symptom, e.g., neck pain] ≤[target]/10 with [functional activity].
2. Pt will improve [objective finding, e.g., cervical rotation] by ≥[measurable target] to allow [activity].
3. Pt will demonstrate ≥[percent]% adherence to [strategy/technique] during [ADL].
4. Pt will perform HEP, transfer, or mobility] with [level of independence] to support [function].

Long-Term Goals (13–25 visits):
1. Pt will increase [strength or ability] by ≥[amount] to support safe [ADL/task].
2. Pt will restore [ROM/ability] to within [target] of normal, enabling [activity].
3. Pt will demonstrate 100% adherence to [technique/precaution] during [ADL/IADL].
4. Pt will independently perform [home program or self-management] to maintain function and prevent recurrence.

ALWAYS use this structure, always begin each statement with 'Pt will', and do NOT add any extra text, dashes, bullets, or lines. Use only info from the above findings.
"""

    result = gpt_call(prompt, max_tokens=400)
    return jsonify({"result": result})

#==================================== PT SAVE NOTES ====================================

@app.route('/pt_eval_builder', methods=['GET', 'POST'])
@login_required
def pt_eval_builder():
    if request.method == 'POST':
        patient_id = request.form.get('patient_id')
        if not patient_id:
            flash('Please assign the note to a patient.', 'danger')
            return render_template('pt_eval_builder.html')

        doc_type = request.form.get('doc_type', 'Evaluation')
        content = request.form.get('generated_note', '')
        fields_json = request.form.get('fields_json', '')

        try:
            note = PTNote(
                patient_id=patient_id,
                content=content,
                user_id=current_user.id,
                doc_type=doc_type,
                fields_json=fields_json
            )
            db.session.add(note)
            db.session.commit()
            flash('Note saved successfully.', 'success')
            return redirect(url_for('view_pt_notes', patient_id=patient_id))
        except Exception as e:
            db.session.rollback()
            flash(f'Error saving note: {str(e)}', 'danger')
            return render_template('pt_eval_builder.html')

    # GET request
    note_id = request.args.get('note_id')
    loaded_note = None
    loaded_doc_type = None
    if note_id:
        note = PTNote.query.filter_by(id=note_id, user_id=current_user.id).first()
        if note:
            if note.fields_json:
                try:
                    loaded_note = json.loads(note.fields_json)
                except Exception:
                    loaded_note = None
            loaded_doc_type = note.doc_type

    return render_template('pt_eval_builder.html', loaded_note=loaded_note, loaded_doc_type=loaded_doc_type)



@app.route('/pt_notes')
@login_required
def pt_notes():
    # Fetch all saved notes for the user (or all, if admin)
    notes = PTNote.query.filter_by(user_id=current_user.id).all()
    return render_template('pt_notes.html', notes=notes)

#==================================== PT EXPORT WORD DOC ====================================
@app.route('/pt_export_word', methods=['POST'])
@login_required
def pt_export_word():
    data = request.get_json()
    buf = pt_export_to_word(data)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='PT_Eval.docx'
    )

def pt_export_to_word(data):
    doc = Document()

    def add_separator():
        doc.add_paragraph('-' * 114)

    # Header line (customize if you want gender/DOB in this header)
    weight = data.get('weight', '')
    height = data.get('height', '')
    bmi = data.get('bmi', '')
    doc.add_paragraph(f"Weight: {weight} lbs       Height: {height}      BMI: {bmi}")
    add_separator()

    doc.add_paragraph(f"Medical Diagnosis: {data.get('meddiag', '')}")
    add_separator()

    doc.add_paragraph("Medical History/HNP:")
    doc.add_paragraph(data.get('history', ''))
    add_separator()

    doc.add_paragraph("Subjective:")
    doc.add_paragraph(data.get('subjective', ''))
    add_separator()

    doc.add_paragraph("Pain:")
    doc.add_paragraph(f"Area/Location of Injury: {data.get('pain_location', '')}")
    doc.add_paragraph(f"Onset/Exacerbation Date: {data.get('pain_onset', '')}")
    doc.add_paragraph(f"Condition of Injury: {data.get('pain_condition', '')}")
    doc.add_paragraph(f"Mechanism of Injury: {data.get('pain_mechanism', '')}")
    doc.add_paragraph(f"Pain Rating (Present/Best/Worst): {data.get('pain_rating', '')}")
    doc.add_paragraph(f"Frequency: {data.get('pain_frequency', '')}")
    doc.add_paragraph(f"Description: {data.get('pain_description', '')}")
    doc.add_paragraph(f"Aggravating Factor: {data.get('pain_aggravating', '')}")
    doc.add_paragraph(f"Relieved By: {data.get('pain_relieved', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Interferes With: {data.get('pain_interferes', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Diagnostic Test(s): {data.get('tests', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"DME/Assistive Device: {data.get('dme', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"PLOF: {data.get('plof', '')}")
    add_separator()

    doc.add_paragraph("Objective:")
    doc.add_paragraph("Posture:")
    doc.add_paragraph(data.get('posture', ''))

    doc.add_paragraph("ROM:")
    doc.add_paragraph(data.get('rom', ''))

    doc.add_paragraph("Muscle Strength Test:")
    doc.add_paragraph(data.get('strength', ''))
    
    doc.add_paragraph("Palpation:")
    doc.add_paragraph(data.get('palpation', ''))

    doc.add_paragraph("Functional Test(s):")
    doc.add_paragraph(data.get('functional', ''))

    doc.add_paragraph("Special Test(s):")
    doc.add_paragraph(data.get('special', ''))

    doc.add_paragraph("Current Functional Mobility Impairment(s):")
    doc.add_paragraph(data.get('impairments', ''))
    add_separator()

    doc.add_paragraph("Assessment Summary:")
    doc.add_paragraph(data.get('summary', ''))
    add_separator()

    doc.add_paragraph("Goals:")
    doc.add_paragraph(data.get('goals', ''))
    add_separator()

    doc.add_paragraph("Frequency:")
    doc.add_paragraph(data.get('frequency', ''))
    add_separator()

    doc.add_paragraph("Intervention:")
    doc.add_paragraph(data.get('intervention', ''))
    add_separator()

    doc.add_paragraph("Treatment Procedures:")
    doc.add_paragraph(data.get('procedures', ''))
    add_separator()

    # --- SOAP ASSESSMENT SECTION ---
    doc.add_paragraph("Soap Assessment:")
    doc.add_paragraph("")

    # Pain section
    pain_lines = [
        "Pain:",
        f"{data.get('pain_location', '')}",
        f"{data.get('pain_rating', '')}",
    ]
    for line in pain_lines:
        doc.add_paragraph(line)
    doc.add_paragraph("")  # blank line after Pain

    # ROM section (split into lines if needed)
    doc.add_paragraph("ROM:")
    for line in data.get('rom', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Palpation section (split into lines if needed)
    doc.add_paragraph("Palpation:")
    for line in data.get('palpation', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Functional Tests section (split into lines if needed)
    doc.add_paragraph("Functional Test(s):")
    for line in data.get('functional', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Goals section (split into lines if needed)
    doc.add_paragraph("Goals:")
    for line in data.get('goals', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

#==================================== PT EXPORT PDF ====================================

@app.route("/pt_export_pdf", methods=["POST"])
@login_required
def pt_export_pdf():
    data = request.get_json()
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    def add_separator():
        nonlocal y
        y -= 8
        c.setLineWidth(0.5)
        c.line(40, y, width - 40, y)
        y -= 16

    def add_paragraph_block(title, text):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in (text or "").split('\n'):
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
        add_separator()

    def add_multiline_section(title, lines):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in lines:
            for subline in line.split('\n'):
                if subline.strip() != "":
                    c.drawString(48, y, subline)
                    y -= 14
                    if y < 60:
                        c.showPage()
                        y = height - 40
            if line == "":
                y -= 8
        add_separator()

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "Physical Therapy Evaluation")
    y -= 30

    # Demographic Header
    gender = data.get('gender', '')
    dob = data.get('dob', '')
    weight = data.get('weight', '')
    height_val = data.get('height', '')
    bmi = data.get('bmi', '')
    bmi_category = data.get('bmi_category', '')
    demographic = f"Gender: {gender}    DOB: {dob}    Weight: {weight} lbs    Height: {height_val}     BMI: {bmi} ({bmi_category})"

    c.setFont("Helvetica", 11)
    c.drawString(40, y, demographic)
    y -= 20
    add_separator()

    # Sections
    add_paragraph_block("Medical Diagnosis:", data.get("meddiag", ""))
    add_paragraph_block("Medical History/HNP:", data.get("history", ""))
    add_paragraph_block("Subjective:", data.get("subjective", ""))

    pain_lines = [
        f"Area/Location of Injury: {data.get('pain_location','')}",
        f"Onset/Exacerbation Date: {data.get('pain_onset','')}",
        f"Condition of Injury: {data.get('pain_condition','')}",
        f"Mechanism of Injury: {data.get('pain_mechanism','')}",
        f"Pain Rating (Present/Best/Worst): {data.get('pain_rating','')}",
        f"Frequency: {data.get('pain_frequency','')}",
        f"Description: {data.get('pain_description','')}",
        f"Aggravating Factor: {data.get('pain_aggravating','')}",
        f"Relieved By: {data.get('pain_relieved','')}",
        f"Interferes With: {data.get('pain_interferes','')}",
        "",
        f"Current Medication(s): {data.get('meds','')}",
        f"Diagnostic Test(s): {data.get('tests','')}",
        f"DME/Assistive Device: {data.get('dme','')}",
        f"PLOF: {data.get('plof','')}",
    ]
    add_multiline_section("Pain:", pain_lines)

    objective_lines = [
        f"Posture:",
        data.get('posture', ''),
        "",
        f"ROM:",
        data.get('rom', ''),
        "",
        f"Muscle Strength Test:",
        data.get('strength', ''),
        "",
        f"Palpation:",
        data.get('palpation', ''),
        "",
        f"Functional Test(s):",
        data.get('functional', ''),
        "",
        f"Special Test(s):",
        data.get('special', ''),
        "",
        f"Current Functional Mobility Impairment(s):",
        data.get('impairments', '')
    ]
    add_multiline_section("Objective:", objective_lines)

    add_paragraph_block("Assessment Summary:", data.get("summary", ""))
    add_paragraph_block("Goals:", data.get("goals", ""))
    add_paragraph_block("Frequency:", data.get("frequency", ""))
    add_paragraph_block("Intervention:", data.get("intervention", ""))
    add_paragraph_block("Treatment Procedures:", data.get("procedures", ""))

    # --- SOAP ASSESSMENT SECTION ---
    soap_lines = []
    soap_lines.append("Soap Assessment:")
    soap_lines.append("")

    # Pain
    soap_lines.append("Pain:")
    soap_lines.append(f"{data.get('pain_location', '')}")
    soap_lines.append(f"{data.get('pain_rating', '')}")
    soap_lines.append("")

    # ROM
    soap_lines.append("ROM:")
    if data.get('rom', ''):
        soap_lines.extend(data.get('rom', '').split('\n'))
    soap_lines.append("")

    # Palpation
    soap_lines.append("Palpation:")
    if data.get('palpation', ''):
        soap_lines.extend(data.get('palpation', '').split('\n'))
    soap_lines.append("")

    # Functional Test(s)
    soap_lines.append("Functional Test(s):")
    if data.get('functional', ''):
        soap_lines.extend(data.get('functional', '').split('\n'))
    soap_lines.append("")

    # Goals
    soap_lines.append("Goals:")
    if data.get('goals', ''):
        soap_lines.extend(data.get('goals', '').split('\n'))
    soap_lines.append("")

    # Output the SOAP section (no separator after)
    c.setFont("Helvetica-Bold", 13)
    c.drawString(40, y, "Soap Assessment:")
    y -= 18
    c.setFont("Helvetica", 11)
    for line in soap_lines[2:]:
        if line == "":
            y -= 8
        else:
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40

    c.save()
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="PT_Eval.pdf",
        mimetype="application/pdf"
    )


#==================================== OT RENDER EVAL ====================================

@app.route("/ot-eval")
@login_required
def ot_eval():
    return render_template("ot_eval.html")
    
#==================================== AI GENERATE OT DIFFERENTIAL DX ====================================

@app.route("/ot_generate_diffdx", methods=["POST"])
@login_required
def ot_generate_diffdx():
    f = request.json.get("fields", {})   # <-- define f here
    dx = f.get("diffdx", "")
    if not dx:
        pain = "; ".join(f"{lbl}: {f.get(key,'')}"
            for lbl, key in [
                ("Area/Location", "pain_location"),
                ("Onset", "pain_onset"),
                ("Condition", "pain_condition"),
                ("Mechanism", "pain_mechanism"),
                ("Rating", "pain_rating"),
                ("Frequency", "pain_frequency"),
                ("Description", "pain_description"),
                ("Aggravating", "pain_aggravating"),
                ("Relieved", "pain_relieved"),
                ("Interferes", "pain_interferes"),
            ])
        dx_prompt = (
            "You are an OT clinical assistant. Based on the following OT evaluation details, "
            "provide a concise statement of the most clinically-associated OT differential diagnosis. "
            "Do NOT state as fact or as a medical diagnosis—use only language such as 'symptoms and clinical findings are associated with or consistent with' the diagnosis. "
            f"Subjective:\n{f.get('subjective','')}\nPain:\n{pain}\n"
        )
        dx = gpt_call(dx_prompt, max_tokens=200)
    return jsonify({"result": dx})

#=================================== AI GENERATE OT SUMMARY ===================================

@app.route("/ot_generate_summary", methods=["POST"])
@login_required
def ot_generate_summary():
    f = request.json.get("fields", {})
    name = (
        f.get("name") or f.get("ot_patient_name") or f.get("patient_name") or f.get("full_name") or "Pt"
    )
    dob = f.get("dob")
    age = "X"
    if dob:
        for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m-%d-%Y"):
            try:
                dob_dt = datetime.strptime(dob, fmt)
                today_dt = date.today()
                age = today_dt.year - dob_dt.year - ((today_dt.month, today_dt.day) < (dob_dt.month, dob_dt.day))
                break
            except Exception:
                continue
    else:
        age = f.get("age", "X")
    gender = f.get("gender", "patient").lower()
    pmh = f.get("history", "no significant history")
    today = f.get("currentdate", date.today().strftime("%m/%d/%Y"))
    subj = f.get("subjective", "")
    moi = f.get("pain_mechanism", "")
    meddiag = f.get("meddiag", "") or f.get("medical_diagnosis", "")
    dx = f.get("diffdx", "")
    # Optional: auto-generate dx if not present
    if not dx:
        pain = "; ".join(f"{lbl}: {f.get(key,'')}"
            for lbl, key in [
                ("Area/Location", "pain_location"),
                ("Onset", "pain_onset"),
                ("Condition", "pain_condition"),
                ("Mechanism", "pain_mechanism"),
                ("Rating", "pain_rating"),
                ("Frequency", "pain_frequency"),
                ("Description", "pain_description"),
                ("Aggravating", "pain_aggravating"),
                ("Relieved", "pain_relieved"),
                ("Interferes", "pain_interferes"),
            ])
        dx_prompt = (
            "You are an OT clinical assistant. Based on the following OT evaluation details, "
            "provide a concise statement of the most clinically-associated OT differential diagnosis. "
            "Do NOT state as fact or as a medical diagnosis—use only language such as 'symptoms and clinical findings are associated with or consistent with' the diagnosis. "
            f"Subjective:\n{subj}\nPain:\n{pain}\n"
        )
        dx = gpt_call(dx_prompt, max_tokens=200)
    strg = f.get("strength", "")
    rom = f.get("rom", "")
    impair = f.get("impairments", "")
    func = f.get("functional", "")

    prompt = (
        "Generate a concise, 7-8 sentence Occupational Therapy assessment summary that is Medicare compliant for OT documentation. "
        "Use only abbreviations (e.g., HEP, ADLs, IADLs, STM, TherEx) and NEVER spell out abbreviations. "
        "Never use 'the patient'; use 'Pt' as the subject. "
        "Do NOT use parentheses, asterisks, or markdown formatting in your response. "
        "Do NOT use 'Diagnosis:' as a label—refer directly to the diagnosis in clinical sentences. "
        "Confirm/Conclude a medical diagnosis—use clinical phrasing such as 'symptoms and clinical findings are associated with' the medical diagnosis and OT clinical impression. "
        f"Start with: \"{name}, a {age} y/o {gender} with relevant history of {pmh}.\" "
        f"Include: OT initial eval on {today} for {subj}. "
        f"If available, mention the mechanism of injury: {moi}. "
        f"State: Pt has symptoms and clinical findings associated with the referring medical diagnosis of {meddiag}. Clinical findings are consistent with OT differential diagnosis of {dx} based on assessment. "
        f"Summarize current impairments (strength: {strg}; ROM: {rom}; balance/mobility: {impair}). "
        f"Summarize functional/activity limitations: {func}. "
        "End with a professional prognosis stating that skilled OT is medically necessary to address impairments and support return to PLOF. "
        "Do NOT use bulleted or numbered lists—compose a single, well-written summary paragraph."
    )
    result = gpt_call(prompt, max_tokens=500)
    return jsonify({"result": result})
    
#==================================== AI GENERATE OT GOALS ====================================

@app.route('/ot_generate_goals', methods=['POST'])
@login_required
def ot_generate_goals():
    fields = request.json.get("fields", {})
    summary = fields.get("summary", "") or "Pt evaluated for functional deficits impacting ADLs/IADLs."
    strength = fields.get("strength", "") or "N/A"
    rom = fields.get("rom", "") or "N/A"
    impairments = fields.get("impairments", "") or "N/A"
    functional = fields.get("functional", "") or "N/A"

    prompt = f"""
You are a clinical assistant helping an occupational therapist write documentation.
Below is a summary of the patient's evaluation and findings:
Summary: {summary}
Strength: {strength}
ROM: {rom}
Impairments: {impairments}
Functional Limitations: {functional}

Using ONLY the above provided eval info, generate clinically-appropriate, Medicare-compliant short-term and long-term OT goals. Focus on ADLs, IADLs, functional participation, use of adaptive equipment, and safety (e.g., dressing, bathing, toileting, home management, transfers, community integration). Each goal must use the "Pt will..." format, specify an activity and level of independence, and be functionally/measurably stated.

ALWAYS follow this exact format—do not add, skip, reorder, or alter any lines or labels.
DO NOT add any explanations, introductions, dashes, bullets, or extra indentation. Output ONLY this structure:

Short-Term Goals (1–12 visits):
1. Pt will [specific ADL/IADL task] with [level of assistance/adaptive strategy] to promote functional independence.
2. Pt will improve [ROM/strength/endurance] by [amount or %] to support [specific functional task].
3. Pt will independently use [adaptive equipment or compensatory technique] during [ADL/IADL].
4. Pt will report pain ≤[target]/10 during [functional task or ADL].

Long-Term Goals (13–25 visits):
1. Pt will complete all [ADL/IADL] independently or with AE as needed.
2. Pt will demonstrate safe performance of [home/community management task] using proper body mechanics and adaptive strategies.
3. Pt will participate in [community activity/home management/transfer] with [level of independence].
4. Pt will maintain functional gains and independently implement all learned safety/adaptive strategies in daily routines.
"""

    result = gpt_call(prompt, max_tokens=400)
    return jsonify({"result": result})

#==================================== OT EXPORT WORD DOC ====================================

@app.route('/ot_export_word', methods=['POST'])
@login_required
def ot_export_word():
    data = request.get_json()
    buf = ot_export_to_word(data)
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='OT_Eval.docx'
    )

def ot_export_to_word(data):
    doc = Document()

    def add_separator():
        doc.add_paragraph('-' * 114)

    # Header
    weight = data.get('weight', '')
    height = data.get('height', '')
    bmi = data.get('bmi', '')
    doc.add_paragraph(f"Weight: {weight} lbs       Height: {height}      BMI: {bmi}")
    add_separator()

    doc.add_paragraph(f"Medical Diagnosis: {data.get('meddiag', '')}")
    add_separator()

    doc.add_paragraph("Medical History/HNP:")
    doc.add_paragraph(data.get('history', ''))
    add_separator()

    doc.add_paragraph("Subjective:")
    doc.add_paragraph(data.get('subjective', ''))
    add_separator()

    doc.add_paragraph("Pain:")
    doc.add_paragraph(f"Area/Location of Injury: {data.get('pain_location', '')}")
    doc.add_paragraph(f"Onset/Exacerbation Date: {data.get('pain_onset', '')}")
    doc.add_paragraph(f"Condition of Injury: {data.get('pain_condition', '')}")
    doc.add_paragraph(f"Mechanism of Injury: {data.get('pain_mechanism', '')}")
    doc.add_paragraph(f"Pain Rating (Present/Best/Worst): {data.get('pain_rating', '')}")
    doc.add_paragraph(f"Frequency: {data.get('pain_frequency', '')}")
    doc.add_paragraph(f"Description: {data.get('pain_description', '')}")
    doc.add_paragraph(f"Aggravating Factor: {data.get('pain_aggravating', '')}")
    doc.add_paragraph(f"Relieved By: {data.get('pain_relieved', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Interferes With: {data.get('pain_interferes', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Diagnostic Test(s): {data.get('tests', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"DME/Assistive Device: {data.get('dme', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(f"PLOF: {data.get('plof', '')}")
    add_separator()

    doc.add_paragraph("Objective:")
    doc.add_paragraph("Posture:")
    doc.add_paragraph(data.get('posture', ''))

    doc.add_paragraph("ROM:")
    doc.add_paragraph(data.get('rom', ''))

    doc.add_paragraph("Muscle Strength Test:")
    doc.add_paragraph(data.get('strength', ''))

    doc.add_paragraph("Palpation:")
    doc.add_paragraph(data.get('palpation', ''))

    doc.add_paragraph("Functional Test(s):")
    doc.add_paragraph(data.get('functional', ''))

    doc.add_paragraph("Special Test(s):")
    doc.add_paragraph(data.get('special', ''))

    doc.add_paragraph("Current Functional Mobility Impairment(s):")
    doc.add_paragraph(data.get('impairments', ''))
    add_separator()

    doc.add_paragraph("Assessment Summary:")
    doc.add_paragraph(data.get('summary', ''))
    add_separator()

    doc.add_paragraph("Goals:")
    doc.add_paragraph(data.get('goals', ''))
    add_separator()

    doc.add_paragraph("Frequency:")
    doc.add_paragraph(data.get('frequency', ''))
    add_separator()

    doc.add_paragraph("Intervention:")
    doc.add_paragraph(data.get('intervention', ''))
    add_separator()

    doc.add_paragraph("Treatment Procedures:")
    doc.add_paragraph(data.get('procedures', ''))
    add_separator()

    # --- SOAP ASSESSMENT SECTION ---
    doc.add_paragraph("Soap Assessment:")
    doc.add_paragraph("")

    # Pain section
    pain_lines = [
        "Pain:",
        f"{data.get('pain_location', '')}",
        f"{data.get('pain_rating', '')}",
    ]
    for line in pain_lines:
        doc.add_paragraph(line)
    doc.add_paragraph("")  # blank line after Pain

    # ROM section (split into lines if needed)
    doc.add_paragraph("ROM:")
    for line in data.get('rom', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Palpation section (split into lines if needed)
    doc.add_paragraph("Palpation:")
    for line in data.get('palpation', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Functional Tests section (split into lines if needed)
    doc.add_paragraph("Functional Test(s):")
    for line in data.get('functional', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    # Goals section (split into lines if needed)
    doc.add_paragraph("Goals:")
    for line in data.get('goals', '').split('\n'):
        doc.add_paragraph(line)
    doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

#==================================== OT EXPORT PDF ====================================

@app.route("/ot_export_pdf", methods=["POST"])
@login_required
def ot_export_pdf():
    data = request.get_json()
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 40

    def add_separator():
        nonlocal y
        y -= 8
        c.setLineWidth(0.5)
        c.line(40, y, width - 40, y)
        y -= 16

    def add_paragraph_block(title, text):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in (text or "").split('\n'):
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
        add_separator()

    def add_multiline_section(title, lines):
        nonlocal y
        c.setFont("Helvetica-Bold", 13)
        c.drawString(40, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for line in lines:
            for subline in line.split('\n'):
                if subline.strip() != "":
                    c.drawString(48, y, subline)
                    y -= 14
                    if y < 60:
                        c.showPage()
                        y = height - 40
            if line == "":
                y -= 8
        add_separator()

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "Occupational Therapy Evaluation")
    y -= 30

    # Demographic Header
    gender = data.get('gender', '')
    dob = data.get('dob', '')
    weight = data.get('weight', '')
    height_val = data.get('height', '')
    bmi = data.get('bmi', '')
    bmi_category = data.get('bmi_category', '')
    demographic = f"Gender: {gender}    DOB: {dob}    Weight: {weight} lbs    Height: {height_val}     BMI: {bmi} ({bmi_category})"

    c.setFont("Helvetica", 11)
    c.drawString(40, y, demographic)
    y -= 20
    add_separator()

    # Sections
    add_paragraph_block("Medical Diagnosis:", data.get("meddiag", ""))
    add_paragraph_block("Medical History/HNP:", data.get("history", ""))
    add_paragraph_block("Subjective:", data.get("subjective", ""))

    pain_lines = [
        f"Area/Location of Injury: {data.get('pain_location','')}",
        f"Onset/Exacerbation Date: {data.get('pain_onset','')}",
        f"Condition of Injury: {data.get('pain_condition','')}",
        f"Mechanism of Injury: {data.get('pain_mechanism','')}",
        f"Pain Rating (Present/Best/Worst): {data.get('pain_rating','')}",
        f"Frequency: {data.get('pain_frequency','')}",
        f"Description: {data.get('pain_description','')}",
        f"Aggravating Factor: {data.get('pain_aggravating','')}",
        f"Relieved By: {data.get('pain_relieved','')}",
        f"Interferes With: {data.get('pain_interferes','')}",
        "",
        f"Current Medication(s): {data.get('meds','')}",
        f"Diagnostic Test(s): {data.get('tests','')}",
        f"DME/Assistive Device: {data.get('dme','')}",
        f"PLOF: {data.get('plof','')}",
    ]
    add_multiline_section("Pain:", pain_lines)

    objective_lines = [
        f"Posture:",
        data.get('posture', ''),
        "",
        f"ROM:",
        data.get('rom', ''),
        "",
        f"Muscle Strength Test:",
        data.get('strength', ''),
        "",
        f"Palpation:",
        data.get('palpation', ''),
        "",
        f"Functional Test(s):",
        data.get('functional', ''),
        "",
        f"Special Test(s):",
        data.get('special', ''),
        "",
        f"Current Functional Mobility Impairment(s):",
        data.get('impairments', '')
    ]
    add_multiline_section("Objective:", objective_lines)

    add_paragraph_block("Assessment Summary:", data.get("summary", ""))
    add_paragraph_block("Goals:", data.get("goals", ""))
    add_paragraph_block("Frequency:", data.get("frequency", ""))
    add_paragraph_block("Intervention:", data.get("intervention", ""))
    add_paragraph_block("Treatment Procedures:", data.get("procedures", ""))

    # --- SOAP ASSESSMENT SECTION ---
    soap_lines = []
    soap_lines.append("Soap Assessment:")
    soap_lines.append("")

    # Pain
    soap_lines.append("Pain:")
    soap_lines.append(f"{data.get('pain_location', '')}")
    soap_lines.append(f"{data.get('pain_rating', '')}")
    soap_lines.append("")

    # ROM
    soap_lines.append("ROM:")
    if data.get('rom', ''):
        soap_lines.extend(data.get('rom', '').split('\n'))
    soap_lines.append("")

    # Palpation
    soap_lines.append("Palpation:")
    if data.get('palpation', ''):
        soap_lines.extend(data.get('palpation', '').split('\n'))
    soap_lines.append("")

    # Functional Test(s)
    soap_lines.append("Functional Test(s):")
    if data.get('functional', ''):
        soap_lines.extend(data.get('functional', '').split('\n'))
    soap_lines.append("")

    # Goals
    soap_lines.append("Goals:")
    if data.get('goals', ''):
        soap_lines.extend(data.get('goals', '').split('\n'))
    soap_lines.append("")

    # Output the SOAP section (no separator after)
    c.setFont("Helvetica-Bold", 13)
    c.drawString(40, y, "Soap Assessment:")
    y -= 18
    c.setFont("Helvetica", 11)
    for line in soap_lines[2:]:
        if line == "":
            y -= 8
        else:
            c.drawString(48, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40

    c.save()
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name="OT_Eval.pdf",
        mimetype="application/pdf"
    )

#==================================== GPT HELPER ====================================

def gpt_call(prompt, max_tokens=700):
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"OpenAI error: {e}"

#==================================== MAIN ====================================
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(debug=True)
